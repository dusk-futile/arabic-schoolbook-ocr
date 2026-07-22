from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from ..schemas import (
    BlockType,
    CanonicalBlock,
    CanonicalDocument,
    Direction,
    FontSizeClass,
    RenderResult,
    TableData,
    TextAlignment,
)
from ..text_runs import segment_mixed_runs
from .docx_validation import validate_docx

_BULLET_PREFIX = re.compile(r"^\s*[\-•●▪◦*]\s*")
_NUMBER_PREFIX = re.compile(r"^\s*\(?[0-9\u0660-\u0669]+[.)\-]\s*")


def _set_rtl_paragraph(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _set_ltr_paragraph(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "0")


def _set_run_font(run: Any, *, rtl: bool) -> None:
    run.font.name = "Arial"
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), "Arial")
    rtl_node = r_pr.find(qn("w:rtl"))
    if rtl_node is None:
        rtl_node = OxmlElement("w:rtl")
        r_pr.append(rtl_node)
    rtl_node.set(qn("w:val"), "1" if rtl else "0")
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:bidi"), "ar-EG")
    lang.set(qn("w:val"), "ar-EG" if rtl else "en-US")


def _add_runs(paragraph: Any, text: str) -> None:
    for run_data in segment_mixed_runs(text):
        run = paragraph.add_run(run_data.text)
        _set_run_font(run, rtl=run_data.direction == Direction.RTL)


def _add_text_with_breaks(paragraph: Any, text: str) -> None:
    lines = text.split("\n")
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run().add_break()
        _add_runs(paragraph, line)


def _configure_section(section: Any) -> None:
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)


def _set_style(
    style: Any, *, size: float, bold: bool, color: str, before: float, after: float
) -> None:
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.15
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attribute}"), "Arial")
    p_pr = style.element.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def _create_styles(document: DocxDocument) -> None:
    # Named override: source_faithful_arabic_schoolbook (A4, Arial, RTL).
    specs = {
        "Normal": (12, False, "111827", 0, 6),
        "Title": (24, True, "0F766E", 0, 12),
        "Heading 1": (20, True, "0F766E", 14, 8),
        "Heading 2": (17, True, "115E59", 12, 7),
        "Heading 3": (14, True, "1F2937", 10, 6),
    }
    for name, (size, bold, color, before, after) in specs.items():
        _set_style(
            document.styles[name],
            size=size,
            bold=bold,
            color=color,
            before=before,
            after=after,
        )
    custom = {
        "OCR Question": (13, True, "1F2937", 8, 4),
        "OCR Answer Option": (12, False, "1F2937", 2, 2),
        "OCR Caption": (10, False, "475569", 4, 8),
        "OCR Note": (11, False, "334155", 6, 6),
        "OCR Unknown": (11, False, "9A3412", 4, 4),
    }
    for name, spec in custom.items():
        style = (
            document.styles[name]
            if name in document.styles
            else document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        )
        _set_style(style, size=spec[0], bold=spec[1], color=spec[2], before=spec[3], after=spec[4])


def _create_numbering(document: DocxDocument, *, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    existing_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "decimal" if ordered else "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1." if ordered else "•")
    level.append(level_text)
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "right")
    level.append(level_justification)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:right"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    existing_num_ids = [
        int(element.get(qn("w:numId"))) for element in numbering.findall(qn("w:num"))
    ]
    num_id = max(existing_num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def _apply_numbering(paragraph: Any, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)


def _paragraph_style(block_type: BlockType) -> str:
    return {
        BlockType.DOCUMENT_TITLE: "Title",
        BlockType.CHAPTER_TITLE: "Heading 1",
        BlockType.HEADING_1: "Heading 1",
        BlockType.HEADING_2: "Heading 2",
        BlockType.HEADING_3: "Heading 3",
        BlockType.QUESTION: "OCR Question",
        BlockType.ANSWER_OPTION: "OCR Answer Option",
        BlockType.CAPTION: "OCR Caption",
        BlockType.DEFINITION_BOX: "OCR Note",
        BlockType.EXAMPLE_BOX: "OCR Note",
        BlockType.NOTE_BOX: "OCR Note",
        BlockType.UNKNOWN: "OCR Unknown",
    }.get(block_type, "Normal")


def _apply_block_formatting(paragraph: Any, block: CanonicalBlock) -> None:
    formatting = block.formatting
    if formatting is None:
        return
    alignment_map = {
        TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if formatting.alignment is not None:
        paragraph.alignment = alignment_map[formatting.alignment]
    if formatting.space_before_pt is not None:
        paragraph.paragraph_format.space_before = Pt(formatting.space_before_pt)
    if formatting.space_after_pt is not None:
        paragraph.paragraph_format.space_after = Pt(formatting.space_after_pt)
    if formatting.left_indent_pt is not None:
        paragraph.paragraph_format.left_indent = Pt(formatting.left_indent_pt)
    if formatting.right_indent_pt is not None:
        paragraph.paragraph_format.right_indent = Pt(formatting.right_indent_pt)
    if formatting.keep_with_next is not None:
        paragraph.paragraph_format.keep_with_next = formatting.keep_with_next
    size_map = {
        FontSizeClass.SMALL: 10,
        FontSizeClass.NORMAL: 12,
        FontSizeClass.MEDIUM: 14,
        FontSizeClass.LARGE: 17,
        FontSizeClass.EXTRA_LARGE: 22,
    }
    for run in paragraph.runs:
        if formatting.bold is not None:
            run.bold = formatting.bold
        if formatting.font_size_class is not None:
            run.font.size = Pt(size_map[formatting.font_size_class])


def _add_block_paragraphs(
    document: DocxDocument,
    block: CanonicalBlock,
    *,
    polished: bool,
    bullet_num_id: int,
    number_num_id: int,
) -> None:
    text = block.output_text(polished)
    if block.block_type in {BlockType.BULLET_LIST, BlockType.NUMBERED_LIST}:
        matcher = _BULLET_PREFIX if block.block_type == BlockType.BULLET_LIST else _NUMBER_PREFIX
        for item in [part for part in text.splitlines() if part.strip()]:
            paragraph = document.add_paragraph(style="Normal")
            if block.paragraph_direction == Direction.LTR:
                _set_ltr_paragraph(paragraph)
            else:
                _set_rtl_paragraph(paragraph)
            _apply_numbering(
                paragraph,
                bullet_num_id if block.block_type == BlockType.BULLET_LIST else number_num_id,
            )
            _add_runs(paragraph, matcher.sub("", item, count=1))
            _apply_block_formatting(paragraph, block)
        return
    paragraphs = text.split("\n\n") if text else [""]
    for paragraph_text in paragraphs:
        paragraph = document.add_paragraph(style=_paragraph_style(block.block_type))
        if block.paragraph_direction == Direction.LTR:
            _set_ltr_paragraph(paragraph)
        else:
            _set_rtl_paragraph(paragraph)
        _add_text_with_breaks(paragraph, paragraph_text)
        _apply_block_formatting(paragraph, block)
        if block.unresolved:
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), "FFF7ED")
            paragraph._p.get_or_add_pPr().append(shading)
        if not paragraph_text and block.block_type in {
            BlockType.UNKNOWN,
            BlockType.EQUATION_IMAGE,
            BlockType.DECORATIVE_REGION,
        }:
            _add_runs(paragraph, f"[{block.block_type.value} - unresolved]")


def _set_table_geometry(table: Any, width_dxa: int) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table_pr = table._tbl.tblPr
    bidi = table_pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        table_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(width_dxa))
    table_width.set(qn("w:type"), "dxa")
    indent = table_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        table_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")


def _add_table(document: DocxDocument, data: TableData) -> None:
    table = document.add_table(rows=data.rows, cols=data.columns)
    table.style = "Table Grid"
    width_dxa = 9893  # A4 width minus 0.7 inch margins.
    column_width = width_dxa // data.columns
    _set_table_geometry(table, width_dxa)
    for column in table.columns:
        column.width = Inches(column_width / 1440)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Inches(column_width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
    for cell_data in data.cells:
        if cell_data.row >= data.rows or cell_data.column >= data.columns:
            continue
        cell = table.cell(cell_data.row, cell_data.column)
        if cell_data.row_span > 1 or cell_data.column_span > 1:
            end_row = min(data.rows - 1, cell_data.row + cell_data.row_span - 1)
            end_column = min(data.columns - 1, cell_data.column + cell_data.column_span - 1)
            cell = cell.merge(table.cell(end_row, end_column))
        paragraph = cell.paragraphs[0]
        _set_rtl_paragraph(paragraph)
        _add_text_with_breaks(paragraph, cell_data.text)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    table_header = OxmlElement("w:tblHeader")
    table_header.set(qn("w:val"), "true")
    header_properties.append(table_header)


def _add_figure(
    document: DocxDocument,
    image_path: Path,
    block: CanonicalBlock,
    *,
    page_width: int,
) -> None:
    """Add a figure at approximately its source-page scale.

    Scaling every crop to the full text width makes narrow portrait crops enormous
    and can force one source page to span several DOCX pages.  The block geometry
    supplies a stable, source-faithful width; the crop dimensions preserve the
    image aspect ratio.  The caps are guardrails for malformed geometry.
    """

    content_width_inches = 6.87  # A4 width minus the configured 0.7 inch margins.
    width_inches = content_width_inches * block.bbox.width / max(page_width, 1)
    width_inches = min(5.8, max(0.6, width_inches))
    with Image.open(image_path) as image:
        image_width, image_height = image.size
    height_inches = width_inches * image_height / max(image_width, 1)
    if height_inches > 7.0:
        scale = 7.0 / height_inches
        width_inches *= scale
        height_inches = 7.0

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(
        str(image_path),
        width=Inches(width_inches),
        height=Inches(height_inches),
    )


def _clear_container(container: Any) -> Any:
    paragraph = container.paragraphs[0]
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)
    return paragraph


def _populate_header_footer(section: Any, blocks: list[CanonicalBlock], *, polished: bool) -> None:
    header_blocks = [block for block in blocks if block.block_type == BlockType.HEADER]
    footer_blocks = [
        block for block in blocks if block.block_type in {BlockType.FOOTER, BlockType.PAGE_NUMBER}
    ]
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    header = _clear_container(section.header)
    footer = _clear_container(section.footer)
    _set_rtl_paragraph(header)
    _set_rtl_paragraph(footer)
    _add_runs(header, " | ".join(block.output_text(polished) for block in header_blocks))
    _add_runs(footer, " | ".join(block.output_text(polished) for block in footer_blocks))


class DocxDocumentRenderer:
    """Semantic Arabic DOCX renderer using the named source-faithful style override."""

    def render_docx(
        self,
        document: CanonicalDocument,
        output_path: Path,
        *,
        polished: bool = False,
    ) -> RenderResult:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.core_properties.title = document.title
        doc.core_properties.author = "Arabic Schoolbook OCR"
        doc.core_properties.comments = "Generated from canonical logical-order Unicode"
        _create_styles(doc)
        bullet_num_id = _create_numbering(doc, ordered=False)
        number_num_id = _create_numbering(doc, ordered=True)
        warnings: list[str] = []
        for page_index, page in enumerate(
            sorted(document.pages, key=lambda item: item.page_number)
        ):
            section = doc.sections[0] if page_index == 0 else doc.add_section(WD_SECTION.NEW_PAGE)
            _configure_section(section)
            _populate_header_footer(section, page.blocks, polished=polished)
            if page.status.value != "COMPLETED":
                paragraph = doc.add_paragraph(style="OCR Unknown")
                _set_rtl_paragraph(paragraph)
                _add_runs(
                    paragraph, f"[Page {page.page_number} failed: {page.error or 'unknown error'}]"
                )
                warnings.append(f"Page {page.page_number} is represented by a failure marker")
                continue
            for block in sorted(page.blocks, key=lambda item: item.reading_order):
                if block.block_type in {BlockType.HEADER, BlockType.FOOTER, BlockType.PAGE_NUMBER}:
                    continue
                if block.block_type == BlockType.FIGURE and block.source_crop:
                    figure_path = output_path.parent.parent / block.source_crop
                    if figure_path.is_file():
                        _add_figure(
                            doc,
                            figure_path,
                            block,
                            page_width=page.width,
                        )
                    else:
                        warnings.append(
                            "Figure crop for page "
                            f"{page.page_number} is missing: {block.source_crop}"
                        )
                        _add_block_paragraphs(
                            doc,
                            block,
                            polished=polished,
                            bullet_num_id=bullet_num_id,
                            number_num_id=number_num_id,
                        )
                elif block.block_type == BlockType.TABLE and block.table is not None:
                    _add_table(doc, block.table)
                else:
                    _add_block_paragraphs(
                        doc,
                        block,
                        polished=polished,
                        bullet_num_id=bullet_num_id,
                        number_num_id=number_num_id,
                    )
        doc.save(str(output_path))
        validation_path = output_path.with_suffix(".validation.json")
        validation = validate_docx(output_path, document, validation_path)
        warnings.extend(validation["warnings"])
        return RenderResult(
            output_path=output_path,
            polished=polished,
            page_count=len(document.pages),
            block_count=sum(len(page.blocks) for page in document.pages),
            warnings=warnings,
            validation_report=validation_path,
        )
