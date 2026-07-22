from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path
from typing import Any

from docx import Document

from ..schemas import BlockType, CanonicalDocument


def _canonical_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _expected_fragments(block_type: BlockType, text: str) -> list[str]:
    if block_type in {BlockType.BULLET_LIST, BlockType.NUMBERED_LIST}:
        prefix = re.compile(r"^\s*(?:[•◦▪‣\-\u2013\u2014]|\d+[.)-])\s*")
        return [
            _canonical_whitespace(prefix.sub("", line, count=1))
            for line in text.splitlines()
            if line.strip()
        ]
    expected = _canonical_whitespace(text)
    return [expected] if expected else []


def _extract_all_text(path: Path) -> str:
    document = Document(str(path))
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return "\n".join(parts)


def validate_docx(
    path: Path, canonical: CanonicalDocument, report_path: Path | None = None
) -> dict[str, Any]:
    required_parts = {"[Content_Types].xml", "word/document.xml", "word/styles.xml"}
    warnings: list[str] = []
    with zipfile.ZipFile(path) as archive:
        corrupt = archive.testzip()
        names = set(archive.namelist())
        missing_parts = sorted(required_parts - names)
        document_xml = archive.read("word/document.xml").decode("utf-8")
    extracted = _extract_all_text(path)
    searchable = _canonical_whitespace(extracted)
    missing_blocks: list[dict[str, Any]] = []
    for page in canonical.pages:
        for block in page.blocks:
            if block.block_type in {BlockType.FIGURE, BlockType.DECORATIVE_REGION}:
                continue
            expected_fragments = _expected_fragments(block.block_type, block.literal_text)
            if block.block_type == BlockType.TABLE and block.table is not None:
                expected_fragments.extend(
                    _canonical_whitespace(cell.text)
                    for cell in block.table.cells
                    if cell.text.strip()
                )
            for expected in expected_fragments:
                if expected and expected not in searchable:
                    missing_blocks.append(
                        {
                            "page": page.page_number,
                            "block_id": block.id,
                            "text": expected[:160],
                        }
                    )
    if corrupt:
        warnings.append(f"Corrupt ZIP member: {corrupt}")
    if missing_parts:
        warnings.append("Required DOCX XML parts are missing")
    if missing_blocks:
        warnings.append(f"{len(missing_blocks)} canonical blocks were not found after reopening")
    report = {
        "docx": path.name,
        "zip_ok": corrupt is None,
        "missing_xml_parts": missing_parts,
        "reopen_ok": True,
        "logical_text_length": len(extracted),
        "section_break_count": document_xml.count('w:type="nextPage"'),
        "bidi_paragraph_count": document_xml.count("<w:bidi"),
        "rtl_run_count": document_xml.count("<w:rtl"),
        "table_count": document_xml.count("<w:tbl>"),
        "missing_blocks": missing_blocks,
        "warnings": warnings,
        "passed": corrupt is None and not missing_parts and not missing_blocks,
    }
    if report_path is not None:
        report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report
