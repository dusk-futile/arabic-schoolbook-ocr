import json
import zipfile
from pathlib import Path

from docx import Document
from PIL import Image

from arabic_schoolbook_ocr.rendering import DocxDocumentRenderer
from arabic_schoolbook_ocr.schemas import (
    BlockType,
    BoundingBox,
    CanonicalBlock,
    CanonicalDocument,
    CanonicalPage,
    TableCell,
    TableData,
)


def test_docx_round_trip_and_rtl_ooxml(tmp_path: Path) -> None:
    text = "تحدث عملية Photosynthesis داخل النبات."
    document = CanonicalDocument(
        title="Fixture",
        source_filename="fixture.pdf",
        source_sha256="0" * 64,
        classification="PUBLIC_DEMO",
        pages=[
            CanonicalPage(
                page_number=1,
                width=1000,
                height=1400,
                source_image="source.png",
                blocks=[
                    CanonicalBlock(
                        block_type=BlockType.HEADING_1,
                        bbox=BoundingBox(x=100, y=100, width=800, height=100),
                        reading_order=0,
                        literal_text="عنوان الدرس",
                        unicode_normalized_text="عنوان الدرس",
                        confidence=0.9,
                    ),
                    CanonicalBlock(
                        block_type=BlockType.BODY_PARAGRAPH,
                        bbox=BoundingBox(x=100, y=250, width=800, height=300),
                        reading_order=1,
                        literal_text=text,
                        unicode_normalized_text=text,
                        confidence=0.9,
                    ),
                ],
            )
        ],
    )
    output = tmp_path / "fixture.docx"
    result = DocxDocumentRenderer().render_docx(document, output)
    report = json.loads(result.validation_report.read_text(encoding="utf-8"))
    assert output.is_file()
    assert report["passed"] is True
    assert report["bidi_paragraph_count"] >= 2
    assert report["rtl_run_count"] >= 2


def test_docx_validation_handles_semantic_lists_and_table_cells(tmp_path: Path) -> None:
    document = CanonicalDocument(
        title="قائمة وجدول",
        source_filename="fixture.png",
        source_sha256="0" * 64,
        classification="PUBLIC_DEMO",
        pages=[
            CanonicalPage(
                page_number=1,
                width=1000,
                height=1400,
                source_image="fixture.png",
                blocks=[
                    CanonicalBlock(
                        block_type=BlockType.BULLET_LIST,
                        bbox=BoundingBox(x=100, y=100, width=800, height=200),
                        reading_order=0,
                        literal_text="• البند الأول\n• البند الثاني",
                        unicode_normalized_text="• البند الأول\n• البند الثاني",
                    ),
                    CanonicalBlock(
                        block_type=BlockType.TABLE,
                        bbox=BoundingBox(x=100, y=400, width=800, height=300),
                        reading_order=1,
                        literal_text="",
                        unicode_normalized_text="",
                        table=TableData(
                            rows=1,
                            columns=2,
                            cells=[
                                TableCell(row=0, column=0, text="الخلية الأولى"),
                                TableCell(row=0, column=1, text="الخلية الثانية"),
                            ],
                        ),
                    ),
                ],
            )
        ],
    )
    output = tmp_path / "lists-and-table.docx"
    result = DocxDocumentRenderer().render_docx(document, output)
    report = json.loads(result.validation_report.read_text(encoding="utf-8"))
    assert report["passed"] is True


def test_docx_renderer_embeds_figure_crop(tmp_path: Path) -> None:
    job = tmp_path / "job"
    crop = job / "pages" / "0001" / "disagreement_crops" / "figure.png"
    crop.parent.mkdir(parents=True)
    Image.new("RGB", (320, 180), "white").save(crop)
    document = CanonicalDocument(
        title="Figure",
        source_filename="fixture.png",
        source_sha256="0" * 64,
        classification="PUBLIC_DEMO",
        pages=[
            CanonicalPage(
                page_number=1,
                width=1000,
                height=1400,
                source_image="pages/0001/source.png",
                blocks=[
                    CanonicalBlock(
                        block_type=BlockType.FIGURE,
                        bbox=BoundingBox(x=100, y=200, width=800, height=450),
                        reading_order=0,
                        literal_text="",
                        unicode_normalized_text="",
                        source_crop="pages/0001/disagreement_crops/figure.png",
                    )
                ],
            )
        ],
    )
    output = job / "output" / "figure.docx"
    result = DocxDocumentRenderer().render_docx(document, output)

    with zipfile.ZipFile(output) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media
    assert not result.warnings


def test_docx_renderer_keeps_narrow_figure_near_source_scale(tmp_path: Path) -> None:
    job = tmp_path / "job"
    crop = job / "pages" / "0001" / "disagreement_crops" / "portrait.png"
    crop.parent.mkdir(parents=True)
    Image.new("RGB", (200, 800), "white").save(crop)
    document = CanonicalDocument(
        title="Portrait figure",
        source_filename="fixture.png",
        source_sha256="0" * 64,
        classification="PUBLIC_DEMO",
        pages=[
            CanonicalPage(
                page_number=1,
                width=2000,
                height=3000,
                source_image="pages/0001/source.png",
                blocks=[
                    CanonicalBlock(
                        block_type=BlockType.FIGURE,
                        bbox=BoundingBox(x=100, y=200, width=200, height=800),
                        reading_order=0,
                        literal_text="",
                        unicode_normalized_text="",
                        source_crop="pages/0001/disagreement_crops/portrait.png",
                    )
                ],
            )
        ],
    )
    output = job / "output" / "portrait.docx"
    DocxDocumentRenderer().render_docx(document, output)

    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    # 0.687 inches at 914400 EMU/inch, far below the old fixed 5.8 inches.
    assert 'cx="628192"' in document_xml


def test_arabic_paragraph_uses_word_wrapping_without_manual_line_breaks(
    tmp_path: Path,
) -> None:
    text = "فقرة عربية طويلة تلتف تلقائيا داخل برنامج وورد دون عودة يدوية بعد كل سطر بصري"
    document = CanonicalDocument(
        title="Wrapping",
        source_filename="fixture.pdf",
        source_sha256="0" * 64,
        classification="PUBLIC_DEMO",
        pages=[
            CanonicalPage(
                page_number=1,
                width=1000,
                height=1400,
                source_image="source.png",
                blocks=[
                    CanonicalBlock(
                        block_type=BlockType.BODY_PARAGRAPH,
                        bbox=BoundingBox(x=100, y=100, width=800, height=200),
                        reading_order=0,
                        literal_text=text,
                        unicode_normalized_text=text,
                    )
                ],
            )
        ],
    )
    output = tmp_path / "wrapping.docx"
    DocxDocumentRenderer().render_docx(document, output)

    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert text in document_xml
    assert "<w:br" not in document_xml


def test_docx_semantics_cover_mixed_text_lists_tables_questions_and_page_breaks(
    tmp_path: Path,
) -> None:
    blocks = [
        CanonicalBlock(
            block_type=BlockType.HEADING_1,
            bbox=BoundingBox(x=100, y=100, width=800, height=80),
            reading_order=0,
            literal_text="عنوان الدرس",
            unicode_normalized_text="عنوان الدرس",
        ),
        CanonicalBlock(
            block_type=BlockType.BODY_PARAGRAPH,
            bbox=BoundingBox(x=100, y=200, width=800, height=180),
            reading_order=1,
            literal_text="النسبة \u0665\u0660\u066a (50%) والسرعة 10 m/s في Chapter 3.",
            unicode_normalized_text=(
                "النسبة \u0665\u0660\u066a (50%) والسرعة 10 m/s في Chapter 3."
            ),
        ),
        CanonicalBlock(
            block_type=BlockType.QUESTION,
            bbox=BoundingBox(x=100, y=400, width=800, height=80),
            reading_order=2,
            literal_text="\u0661. ما الإجابة الصحيحة؟",
            unicode_normalized_text="\u0661. ما الإجابة الصحيحة؟",
        ),
        CanonicalBlock(
            block_type=BlockType.ANSWER_OPTION,
            bbox=BoundingBox(x=100, y=490, width=800, height=60),
            reading_order=3,
            literal_text="(أ) الاختيار الأول",
            unicode_normalized_text="(أ) الاختيار الأول",
        ),
        CanonicalBlock(
            block_type=BlockType.BULLET_LIST,
            bbox=BoundingBox(x=100, y=560, width=800, height=100),
            reading_order=4,
            literal_text="• بند أول\n• بند ثان",
            unicode_normalized_text="• بند أول\n• بند ثان",
        ),
        CanonicalBlock(
            block_type=BlockType.NUMBERED_LIST,
            bbox=BoundingBox(x=100, y=670, width=800, height=100),
            reading_order=5,
            literal_text="\u0661. خطوة أولى\n\u0662. خطوة ثانية",
            unicode_normalized_text="\u0661. خطوة أولى\n\u0662. خطوة ثانية",
        ),
        CanonicalBlock(
            block_type=BlockType.TABLE,
            bbox=BoundingBox(x=100, y=780, width=800, height=300),
            reading_order=6,
            literal_text="",
            unicode_normalized_text="",
            table=TableData(
                rows=2,
                columns=2,
                cells=[
                    TableCell(row=0, column=0, text="العنوان"),
                    TableCell(row=0, column=1, text="القيمة"),
                    TableCell(row=1, column=0, text="سطر أول\nسطر ثان"),
                    TableCell(row=1, column=1, text="25%"),
                ],
            ),
        ),
    ]
    document = CanonicalDocument(
        title="Semantics",
        source_filename="fixture.pdf",
        source_sha256="0" * 64,
        classification="PUBLIC_DEMO",
        pages=[
            CanonicalPage(
                page_number=1,
                width=1000,
                height=1400,
                source_image="page1.png",
                blocks=blocks,
            ),
            CanonicalPage(
                page_number=2,
                width=1000,
                height=1400,
                source_image="page2.png",
                blocks=[
                    CanonicalBlock(
                        block_type=BlockType.BODY_PARAGRAPH,
                        bbox=BoundingBox(x=100, y=100, width=800, height=100),
                        reading_order=0,
                        literal_text="صفحة ثانية",
                        unicode_normalized_text="صفحة ثانية",
                    )
                ],
            ),
        ],
    )
    output = tmp_path / "semantic-edge-cases.docx"
    result = DocxDocumentRenderer().render_docx(document, output)
    reopened = Document(output)

    assert len(reopened.sections) == 2
    assert any(paragraph.style.name == "Heading 1" for paragraph in reopened.paragraphs)
    assert len(reopened.tables) == 1
    assert reopened.tables[0].cell(1, 0).text == "سطر أول\nسطر ثان"
    with zipfile.ZipFile(output) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
    assert document_xml.count("<w:numPr>") >= 4
    paragraph_text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
    assert "m/s" in paragraph_text
    assert "Chapter 3" in paragraph_text
    assert "\u0665\u0660\u066a" in paragraph_text and "50%" in paragraph_text
    assert result.validation_report is not None
