"""Word output with real right-to-left semantics.

python-docx has no API for bidi, and getting it wrong is not cosmetic: a
paragraph without <w:bidi/> reorders punctuation to the wrong end of the line,
and a run without <w:rtl/> plus a complex-script font mapping renders as boxes
on machines that lack the fallback. Both are set explicitly here.

Styles are semantic (Heading 1/2, Normal, List Paragraph) because Braille is a
linear medium - the embosser needs to know "this is a heading", not where the
heading sat on the page.
"""
from __future__ import annotations

import re
from typing import Iterable, List, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt

from .model import (BODY, CAPTION, FOOTNOTE, HEADING, LIST_ITEM, SUBHEADING,
                    TITLE, Para)

ARABIC_FONT = "Simplified Arabic"
FALLBACK_FONTS = ["Traditional Arabic", "Arial", "Times New Roman"]


# Whitespace that must never reach an embosser. A Braille cell is literal: a
# no-break space is a real cell, a tab is undefined, a line separator becomes a
# hard line ending mid-paragraph, and a bidi control renders as a stray cell or
# silently reorders the line. The only whitespace allowed inside a paragraph is
# a single U+0020; the only paragraph separator is the Word paragraph mark
# itself, never a character inside the text.
_KILL = dict.fromkeys(map(ord, (
    "\u00a0\u2007\u202f\u2009\u200a\u2028\u2029\u000b\u000c"
    "\u200b\u200c\u200d\ufeff\u00ad\u2060"
    "\u200e\u200f\u202a\u202b\u202c\u202d\u202e"
    "\u2066\u2067\u2068\u2069\u061c"
)), " ")


def emboss_safe(text: str) -> str:
    """Collapse a paragraph to exactly one line with single spaces.

    Applied to every paragraph on the way into the document, so the guarantee
    holds by construction rather than by whatever the upstream stages happened
    to produce. eval/audit_docx.py re-checks it on the finished file.
    """
    if not text:
        return ""
    text = text.translate(_KILL)
    text = text.replace("\t", " ")
    # a newline inside a paragraph is a soft wrap that was never resolved
    text = re.sub(r"[\r\n]+", " ", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def _el(tag: str, **attrs) -> OxmlElement:
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), v)
    return e


def _set_rtl_run(run, font_name: str, size_pt: float):
    rPr = run._element.get_or_add_rPr()
    rPr.append(_el("w:rtl"))
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = _el("w:rFonts")
        rPr.insert(0, rFonts)
    # cs = "complex script": the attribute Word actually uses for Arabic
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    szCs = _el("w:szCs")
    szCs.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(szCs)


def _set_rtl_para(par, justify: bool = True):
    pPr = par._element.get_or_add_pPr()
    pPr.append(_el("w:bidi"))
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.RIGHT


def _set_section_rtl(section):
    sectPr = section._sectPr
    sectPr.append(_el("w:bidi"))


STYLE_FOR = {
    TITLE: ("Title", 20, False),
    HEADING: ("Heading 1", 16, False),
    SUBHEADING: ("Heading 2", 14, False),
    BODY: ("Normal", 13, True),
    LIST_ITEM: ("List Paragraph", 13, False),
    FOOTNOTE: ("Normal", 10, False),
    CAPTION: ("Normal", 11, False),
}


def build_docx(paras: Iterable[Para], out_path: str,
               font: str = ARABIC_FONT,
               mark_uncertain: bool = True,
               title: Optional[str] = None) -> str:
    doc = Document()
    for section in doc.sections:
        _set_section_rtl(section)

    normal = doc.styles["Normal"]
    normal.font.name = font
    normal.font.size = Pt(13)

    if title:
        p = doc.add_paragraph(emboss_safe(title), style="Title")
        _set_rtl_para(p, justify=False)
        for r in p.runs:
            _set_rtl_run(r, font, 20)

    for para in paras:
        text = emboss_safe(para.text)
        if not text:
            continue
        style, size, justify = STYLE_FOR.get(para.kind, ("Normal", 13, True))
        try:
            p = doc.add_paragraph(style=style)
        except KeyError:
            p = doc.add_paragraph()
        run = p.add_run(text)
        _set_rtl_para(p, justify=justify)
        _set_rtl_run(run, font, size)
        if mark_uncertain and (para.flags or para.conf < 0.75):
            # A highlight is the fastest possible affordance for a proofreader:
            # it turns "read the whole book" into "look at the yellow bits".
            hl = _el("w:highlight")
            hl.set(qn("w:val"), "yellow")
            run._element.get_or_add_rPr().append(hl)
    doc.save(out_path)
    return out_path


def build_plain_text(paras: Iterable[Para], out_path: str) -> str:
    """Canonical intermediate: one paragraph per block, blank line between.

    A single newline never appears inside a paragraph. That single rule is what
    makes the file safe to feed to liblouis or an embosser.
    """
    chunks: List[str] = []
    for p in paras:
        t = emboss_safe(p.text)
        if t:
            chunks.append(t)
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(chunks) + "\n")
    return out_path
